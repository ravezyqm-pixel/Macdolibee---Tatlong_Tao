from pathlib import Path
import subprocess
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "documentation"
OUT_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUT_DIR / "MacdoLibee_Final_Documentation.docx"
ERD_PATH = OUT_DIR / "macdolibee_erd.png"
WIREFRAME_PATH = OUT_DIR / "macdolibee_wireframe.png"


def font(size, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def run_mysql(sql):
    result = subprocess.run(
        ["mysql", "-uroot", "--batch", "--raw", "--skip-column-names", "macdolibee2", "-e", sql],
        cwd=ROOT,
        check=True,
        capture_output=True,
        text=True,
    )
    return result.stdout.strip()


def draw_wrapped(draw, xy, text, width, fnt, fill="#111111", line_gap=6):
    x, y = xy
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = (current + " " + word).strip()
        bbox = draw.textbbox((0, 0), trial, font=fnt)
        if bbox[2] - bbox[0] <= width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap
    return y


def draw_table_entity(draw, x, y, w, title, rows, fill="#fff8e8", outline="#bc1c24"):
    title_font = font(24, True)
    row_font = font(17)
    row_h = 30
    h = 48 + row_h * len(rows) + 14
    draw.rounded_rectangle([x, y, x + w, y + h], radius=14, fill=fill, outline=outline, width=4)
    draw.rounded_rectangle([x, y, x + w, y + 46], radius=14, fill=outline)
    draw.text((x + 16, y + 10), title, font=title_font, fill="#ffffff")
    yy = y + 58
    for row in rows:
        draw.text((x + 16, yy), row, font=row_font, fill="#111111")
        yy += row_h
    return (x, y, x + w, y + h)


def line_with_label(draw, p1, p2, label, label_offset=(0, 0)):
    draw.line([p1, p2], fill="#333333", width=4)
    mid = ((p1[0] + p2[0]) // 2 + label_offset[0], (p1[1] + p2[1]) // 2 + label_offset[1])
    label_font = font(18, True)
    bbox = draw.textbbox((0, 0), label, font=label_font)
    pad = 6
    draw.rounded_rectangle(
        [mid[0] - pad, mid[1] - pad, mid[0] + bbox[2] + pad, mid[1] + bbox[3] + pad],
        radius=8,
        fill="#ffffff",
        outline="#888888",
    )
    draw.text(mid, label, font=label_font, fill="#111111")


def make_erd():
    img = Image.new("RGB", (1800, 1200), "#fffaf0")
    draw = ImageDraw.Draw(img)
    title_font = font(42, True)
    subtitle_font = font(22)
    draw.text((60, 38), "MacdoLibee Database ERD", font=title_font, fill="#bc1c24")
    draw.text((60, 92), "Crow's-foot style logical view with implemented and snapshot relationships", font=subtitle_font, fill="#333333")

    foods = draw_table_entity(
        draw,
        80,
        190,
        450,
        "FOODS",
        [
            "PK id",
            "name",
            "category",
            "price",
            "price_w_vat",
            "image_path",
            "created_at",
        ],
    )
    food_ingredients = draw_table_entity(
        draw,
        680,
        215,
        500,
        "FOOD_INGREDIENTS",
        [
            "PK id",
            "FK food_name -> foods.name",
            "FK ingredient_name -> ingredients.name",
            "quantity",
            "UNIQUE food_name + ingredient_name",
        ],
        fill="#fff3d0",
    )
    ingredients = draw_table_entity(
        draw,
        1320,
        190,
        400,
        "INGREDIENTS",
        ["PK id", "name", "cost", "created_at"],
    )
    orders = draw_table_entity(
        draw,
        80,
        700,
        520,
        "ORDERS",
        [
            "PK id",
            "unit_count",
            "sub_total",
            "product_list",
            "barcode",
            "order_time",
            "completed_at",
            "status",
        ],
        fill="#f6f8ff",
        outline="#176b87",
    )
    users = draw_table_entity(
        draw,
        750,
        720,
        400,
        "USERS",
        ["PK id", "username", "password", "role", "created_at"],
        fill="#f5fff0",
        outline="#4c8d35",
    )
    electricity = draw_table_entity(
        draw,
        1280,
        730,
        420,
        "ELECTRICITY",
        ["PK id", "name", "price", "created_at"],
        fill="#f8f0ff",
        outline="#6d3fb2",
    )

    line_with_label(draw, (foods[2], foods[1] + 140), (food_ingredients[0], food_ingredients[1] + 140), "1 to 0..N")
    line_with_label(draw, (food_ingredients[2], food_ingredients[1] + 140), (ingredients[0], ingredients[1] + 120), "0..N to 1")
    line_with_label(draw, (foods[0] + 250, foods[3]), (orders[0] + 260, orders[1]), "snapshot product_list", (-110, 20))

    note_font = font(18)
    notes = [
        "Implemented relationship: FOODS and INGREDIENTS are many-to-many through FOOD_INGREDIENTS.",
        "ORDERS stores product_list as a transaction snapshot, so order lines are not foreign-key enforced.",
        "USERS controls admin/employee access but is not currently linked to ORDERS by foreign key.",
        "ELECTRICITY is an independent cost reference table in the current database.",
    ]
    yy = 1050
    for note in notes:
        draw.text((80, yy), "- " + note, font=note_font, fill="#333333")
        yy += 30

    img.save(ERD_PATH, "PNG")


def wire_box(draw, x, y, w, h, title, body, fill="#ffffff", outline="#bc1c24"):
    title_font = font(24, True)
    body_font = font(17)
    draw.rounded_rectangle([x, y, x + w, y + h], radius=18, fill=fill, outline=outline, width=4)
    draw.rounded_rectangle([x, y, x + w, y + 48], radius=18, fill=outline)
    draw.text((x + 18, y + 12), title, font=title_font, fill="#ffffff")
    yy = y + 64
    for line in body:
        yy = draw_wrapped(draw, (x + 18, yy), line, w - 36, body_font, "#111111", 4) + 4
    return (x, y, x + w, y + h)


def arrow(draw, p1, p2):
    draw.line([p1, p2], fill="#444444", width=4)
    x1, y1 = p1
    x2, y2 = p2
    if x2 >= x1:
        points = [(x2, y2), (x2 - 18, y2 - 10), (x2 - 18, y2 + 10)]
    else:
        points = [(x2, y2), (x2 + 18, y2 - 10), (x2 + 18, y2 + 10)]
    draw.polygon(points, fill="#444444")


def make_wireframe():
    img = Image.new("RGB", (1800, 1300), "#fffaf0")
    draw = ImageDraw.Draw(img)
    title_font = font(42, True)
    draw.text((60, 38), "MacdoLibee System UI / Wireframe", font=title_font, fill="#bc1c24")

    login = wire_box(
        draw,
        70,
        140,
        390,
        250,
        "Login",
        ["Username + password", "Role check: Admin or Employee", "Customer button opens ordering flow"],
    )
    customer = wire_box(
        draw,
        555,
        100,
        500,
        300,
        "Customer Ordering",
        ["Choose Dine In or Take Out", "Browse categories and food cards", "Add/minus quantity, view cart summary", "Place order and generate barcode"],
        fill="#fff7d6",
    )
    admin = wire_box(
        draw,
        555,
        500,
        500,
        320,
        "Admin / Employee Panel",
        ["Food category container at left", "Order data table at center", "Complete task action for pending orders", "Tools panel: Dashboard and Restaurant Stat are admin-only"],
        fill="#fff3d0",
    )
    dashboard = wire_box(
        draw,
        1180,
        120,
        530,
        320,
        "Dashboard",
        ["Daily, weekly, monthly order counts", "Top 5 best foods table", "Users/Admins table with Add, Edit, Delete"],
        fill="#f5fff0",
        outline="#4c8d35",
    )
    stats = wire_box(
        draw,
        1180,
        520,
        530,
        320,
        "Restaurant Stat",
        ["Food price list with VAT", "Ingredient cost list with VAT", "Add, Edit, Delete foods and ingredients"],
        fill="#f6f8ff",
        outline="#176b87",
    )
    viewer = wire_box(
        draw,
        410,
        930,
        980,
        250,
        "Admin Food Category Viewer",
        ["Shows food cards for selected category", "Each food card displays image, name, and price", "Add, Edit, Delete controls manage foods by category"],
        fill="#ffffff",
    )

    arrow(draw, (login[2], 250), (customer[0], 250))
    arrow(draw, (login[2], 320), (admin[0], 640))
    arrow(draw, (admin[2], 595), (dashboard[0], 280))
    arrow(draw, (admin[2], 695), (stats[0], 680))
    arrow(draw, (admin[0] + 150, admin[3]), (viewer[0] + 300, viewer[1]))

    draw.text((60, 1220), "Design concept: fast counter-service workflow, visible role restrictions, and card-based menu management.", font=font(22, True), fill="#333333")
    img.save(WIREFRAME_PATH, "PNG")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "DADCE0")
        borders.append(tag)


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], "F4C125")
        header_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            set_cell_width(header_cells[i], widths[i])
        for p in header_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cells[i], widths[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.line_spacing = 1.2
                p.paragraph_format.space_after = Pt(3)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.add_run(text)
    return p


def create_doc():
    make_erd()
    make_wireframe()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(0)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 2

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    r = title.add_run("MACDOLIBEE FOOD ORDERING AND RESTAURANT MANAGEMENT SYSTEM")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(16)

    for line in [
        "Final Project Documentation",
        "",
        "Prepared by: [Your Name / Group Members]",
        "Date of Submission: May 16, 2026",
        "",
        "Statement of Compliance:",
        "This documentation is prepared in compliance with the required final project format, including the ERD, relationship degree/connectivity/cardinality, UI wireframe, analysis, recommendation, Arial font, double spacing, and site references.",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    add_heading(doc, "II. Statement of the Problem")
    for text in [
        "MacdoLibee is a food ordering and restaurant management system designed for a fast-service restaurant environment. The problem addressed by the project is the difficulty of managing customer orders, food records, ingredient and cost information, sales summaries, and role-based staff access when these tasks are handled manually or through disconnected tools.",
        "Without a centralized system, the restaurant may experience slow order processing, incomplete transaction records, duplicated food information, inconsistent price and VAT computation, unclear employee permissions, and limited visibility into daily, weekly, and monthly performance. The project therefore asks: How can a local restaurant system organize customer ordering, admin monitoring, employee task completion, food maintenance, and dashboard reporting in one database-backed desktop application?",
        "The final system responds to this problem through a Java Swing interface connected to a MySQL database. Customers can place orders through category-based food cards, staff can complete pending orders, and administrators can view dashboards, maintain users, manage foods and ingredients, and inspect restaurant statistics.",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "III. Decision Criteria and Alternative Solutions")
    add_paragraph(doc, "The project alternatives were evaluated according to the following criteria: cost, ease of use, development time, database reliability, maintainability, role-based access, reporting capability, and fit for a school desktop project.")
    add_table(
        doc,
        ["Alternative", "Advantages", "Limitations", "Evaluation"],
        [
            ["Manual logbook", "Lowest setup cost; easy to start.", "Slow searching, weak reporting, high risk of missing records.", "Not recommended because it cannot support dashboards or reliable data history."],
            ["Spreadsheet tracker", "Familiar and inexpensive; simple formulas.", "Weak multi-user control, limited validation, poor transaction workflow.", "Useful for prototypes but not enough for live ordering and role control."],
            ["Web-based system", "Accessible through browsers and easier to deploy to many devices.", "Requires more setup, hosting, web security, and browser compatibility work.", "Good future option, but too broad for the current local desktop scope."],
            ["Java Swing + MySQL desktop system", "Fits NetBeans workflow, supports local database, provides visual forms, and separates admin/employee features.", "Requires local Java and MySQL setup; UI is desktop-only.", "Selected because it balances feasibility, database reliability, and project requirements."],
        ],
        [Inches(1.35), Inches(1.75), Inches(1.85), Inches(1.55)],
    )
    add_paragraph(doc, "Based on these criteria, the selected solution is a Java Swing desktop application using MySQL. This option offers a realistic balance between usability, technical learning, and database-backed reliability.")

    add_heading(doc, "IV. Hardware and Software Specification")
    add_heading(doc, "Hardware Requirements", 2)
    add_table(
        doc,
        ["Component", "Minimum Specification", "Recommended Specification"],
        [
            ["Processor", "64-bit dual-core processor", "Intel Core i3 / Ryzen 3 or higher"],
            ["Memory", "4 GB RAM", "8 GB RAM or higher"],
            ["Storage", "500 MB free space for project and database", "2 GB free space for database, images, and backups"],
            ["Display", "1366 x 768 resolution", "1920 x 1080 resolution for dashboard and admin panels"],
            ["Peripheral", "Keyboard and mouse", "Optional barcode scanner or receipt printer for future deployment"],
        ],
        [Inches(1.4), Inches(2.4), Inches(2.4)],
    )
    add_heading(doc, "Software Requirements", 2)
    add_table(
        doc,
        ["Software", "Purpose"],
        [
            ["Java SE / JDK 25-compatible runtime", "Runs and compiles the Java Swing desktop application."],
            ["Apache NetBeans IDE", "Used for project development, frame editing, compilation, and execution."],
            ["MySQL Server 8.0", "Stores users, foods, ingredients, orders, VAT, image paths, and dashboard data."],
            ["MySQL Connector/J 9.6.0", "Provides JDBC connectivity between Java and MySQL."],
            ["Windows operating system", "Development and execution environment used for the project."],
            ["Java Swing and AWT", "Provides the desktop graphical user interface."],
        ],
        [Inches(2.2), Inches(4.3)],
    )

    add_heading(doc, "D. Entity Relationship Diagram for the Database Design")
    doc.add_picture(str(ERD_PATH), width=Inches(6.5))
    cap = doc.add_paragraph("Figure 1. ERD of the MacdoLibee database design.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Database Entities", 2)
    add_table(
        doc,
        ["Entity", "Primary Purpose", "Key Fields"],
        [
            ["users", "Stores staff login accounts and roles.", "id, username, password, role, created_at"],
            ["foods", "Stores food menu records, prices, VAT values, categories, and image paths.", "id, name, category, price, price_w_vat, image_path"],
            ["ingredients", "Stores ingredient names and costs.", "id, name, cost, created_at"],
            ["food_ingredients", "Associative table between foods and ingredients.", "id, food_name, ingredient_name, quantity"],
            ["orders", "Stores order transaction snapshots, barcode, status, and completion time.", "id, unit_count, sub_total, product_list, barcode, order_time, completed_at, status"],
            ["electricity", "Stores utility cost references.", "id, name, price, created_at"],
        ],
        [Inches(1.35), Inches(2.65), Inches(2.5)],
    )

    add_heading(doc, "Relationship Degree, Connectivity, and Cardinality", 2)
    add_table(
        doc,
        ["Relationship", "Degree", "Connectivity", "Cardinality / Participation"],
        [
            ["foods to food_ingredients", "Binary", "One-to-many", "One food may appear in zero or many food_ingredients rows; each food_ingredients row references exactly one food."],
            ["ingredients to food_ingredients", "Binary", "One-to-many", "One ingredient may appear in zero or many food_ingredients rows; each food_ingredients row references exactly one ingredient."],
            ["foods to ingredients through food_ingredients", "Binary relationship resolved by associative entity", "Many-to-many", "A food may use many ingredients, and an ingredient may be used by many foods. The quantity attribute belongs to the associative table."],
            ["orders to foods", "Conceptual binary, not FK-enforced", "Many-to-many snapshot", "An order can contain many foods and a food can appear in many orders, but the current database stores this as product_list text for transaction history."],
            ["users to admin/employee access", "Unary classification attribute", "One user has one role", "Each user row has one role value: admin or employee. Role determines dashboard/stat access."],
            ["electricity", "No relationship in current schema", "Standalone reference", "The table is independent and can be connected to cost analysis in future versions."],
        ],
        [Inches(1.6), Inches(1.2), Inches(1.35), Inches(2.35)],
    )

    add_heading(doc, "E. UI / Wireframe for the System Project")
    doc.add_picture(str(WIREFRAME_PATH), width=Inches(6.5))
    cap = doc.add_paragraph("Figure 2. UI wireframe and navigation concept for the MacdoLibee system.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "The wireframe uses a counter-service restaurant design concept. The customer side emphasizes simple category browsing, food images, quantity controls, and order confirmation. The staff side emphasizes order completion and record visibility. The administrator side adds dashboard reporting, user management, food CRUD, ingredient CRUD, VAT visibility, and restaurant statistics.")

    add_heading(doc, "V. Analysis")
    analysis_paragraphs = [
        "The MacdoLibee system was developed as a desktop-based ordering and restaurant management application. Its main value is that it connects the customer ordering process, administrative monitoring, and database maintenance into one workflow. The customer does not need to interact with the database directly; instead, the system presents food categories and food cards that hide the complexity of data storage. This makes the ordering process faster and easier to understand.",
        "The system uses a MySQL database as the central storage layer. This is important because the restaurant data is not limited to temporary screen values. Users, foods, ingredients, orders, status values, VAT fields, and image paths remain available after the application is closed. Compared with manual records, this improves data continuity and allows the system to generate summaries such as daily, weekly, and monthly order counts.",
        "A significant design decision is the use of role-based login. Admin accounts can open the Dashboard and Restaurant Stat pages, while employee accounts can enter the order panel in a restricted mode. This improves operational control because employees can complete tasks and view orders without being able to manage dashboard records, edit restaurant statistics, or change administrative data. The restriction reduces the chance of accidental data changes.",
        "The ordering workflow relies on a cart utility that stores food selections and quantities while the customer is building an order. When the order is placed, the system writes an order record containing unit count, subtotal, product list, barcode, timestamp, and status. The status begins as on-going and can later be changed to completed by authorized staff. This creates a practical workflow for restaurants where orders must move from pending to fulfilled.",
        "The database design contains normalized parts and snapshot parts. The food and ingredient relationship is normalized through the food_ingredients table. This is appropriate because a food item may use several ingredients and one ingredient may be used by several food items. The associative table also stores quantity, which belongs to the relationship rather than to either entity alone.",
        "The order record, however, uses product_list as a text snapshot. This choice is simpler for a school desktop project because it preserves the food names and quantities at the time of purchase without requiring an order_items table. The advantage is implementation speed and readable transaction history. The limitation is that reports about individual food sales require parsing text rather than joining normalized order line records. This was handled in the dashboard by parsing product_list to compute top-selling foods.",
        "The dashboard provides decision support for administrators. Daily, weekly, and monthly order counts show activity trends. The top five best foods list helps identify what products are selling. The users and admins table allows management of staff accounts. These features convert raw database rows into information that supports restaurant decisions.",
        "The Restaurant Stat module supports maintenance of food and ingredient data. Administrators can add, edit, and delete food and ingredient records. The system also calculates VAT at 12 percent using the price_w_vat field. This reduces manual computation errors and makes prices clearer in reports. Updating VAT automatically during add/edit operations also keeps future data consistent.",
        "Food images improve the usability of the ordering interface. The project stores image_path values in the foods table and uses project-local PNG files named after food IDs. This design is simple because image records can be filled automatically as src/images/<id>.png. If an image is missing, the system can fall back to a no-image label, preventing the interface from breaking.",
        "The admin food category viewer was expanded to allow add, edit, and delete operations directly inside each category. This improves workflow because the administrator no longer needs to leave the category view to manage menu items. Selecting a food card before editing or deleting also matches the visual style of the system, which is based on food cards rather than only tables.",
        "From a user interface perspective, the project uses large buttons, category panels, tables, food cards, and clear section headers. These choices support quick recognition and reduce the number of steps required to perform common tasks. The color style uses red, yellow, cream, and white to match a fast-food brand identity while keeping controls visually consistent.",
        "The main technical risk is that some relationships are not fully normalized. Orders are stored with a product_list snapshot rather than an order_items table, and users are not linked to completed orders. This does not stop the current system from working, but it limits future reporting such as employee performance, item-level sales by date, and refund handling. A future version should introduce order_items and user_id fields to improve relational accuracy.",
        "Security is another area for improvement. The current users table stores passwords as plain text. This is acceptable only for a classroom prototype, not for production. A real deployment should hash passwords, enforce password rules, and prevent direct exposure of credentials. Employee and admin permissions should also be enforced consistently in every frame, not only through visible button restrictions.",
        "Overall, the project demonstrates a complete transaction-centered restaurant system. It includes customer ordering, staff order completion, administrative reporting, CRUD operations, VAT calculation, food image support, and MySQL persistence. The outcome is a usable local desktop application that can support restaurant-style operations while still leaving clear room for future normalization and security improvements.",
    ]
    for text in analysis_paragraphs:
        add_paragraph(doc, text)

    add_heading(doc, "VI. Recommendation")
    add_paragraph(doc, "The recommended course of action is to continue using the Java Swing and MySQL solution for the final project because it already satisfies the required ordering, dashboard, CRUD, and role-based management features. The system should be presented as a local desktop restaurant management prototype with a clear explanation of its current limits and future improvement path.")
    add_heading(doc, "Lessons Learned", 2)
    add_bullets(
        doc,
        [
            "A database-backed system is more reliable than manual records because order, user, food, and ingredient data remain persistent.",
            "Role-based access is necessary because admin tasks and employee tasks should not have the same permissions.",
            "Dashboard summaries are useful only when the underlying data is consistently updated, especially order status and completion time.",
            "VAT and image paths should be automated to avoid repetitive manual updates.",
            "A simple interface can still be powerful when screens are organized around real workflows such as ordering, completing orders, and maintaining menus.",
        ],
    )
    add_heading(doc, "Recommended Plan of Action", 2)
    add_numbered(
        doc,
        [
            "Add an order_items table so food-level sales can be queried through joins instead of parsing product_list text.",
            "Add user_id or completed_by fields to orders so employee activity can be tracked accurately.",
            "Replace plain-text passwords with hashed passwords before any real deployment.",
            "Add backup and restore procedures for the MySQL database.",
            "Improve image management by adding an image upload button when adding or editing foods.",
            "Continue testing the system using admin and employee accounts to verify that permission restrictions are consistently applied.",
        ],
    )

    add_heading(doc, "VII. Site References")
    refs = [
        "Oracle. Java SE Documentation. https://www.oracle.com/java/technologies/java-se-doc.html",
        "MySQL. MySQL 8.0 Reference Manual: Data Types. https://dev.mysql.com/doc/refman/8.0/en/data-types.html",
        "Apache NetBeans. General Java Development Learning Trail. https://netbeans.apache.org/tutorial/main/kb/docs/java-se/",
        "Microsoft Learn. User Interface Principles for Windows apps. https://learn.microsoft.com/en-us/windows/win32/appuistart/-user-interface-principles",
        "Digital.gov. Wireframes. https://digital.gov/guides/research-collaboration/designing/wireframing",
        "Runestone Academy. Entity-relationship diagrams. https://runestone.academy/ns/books/published/practical_db/PART2_DATA_MODELING/02-ERD/ERD.html",
        "University of Pittsburgh. Entity-Relationship Model notes. https://people.cs.pitt.edu/~chang/156/03ERmodel.html",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Bullet")

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    create_doc()
    print(DOCX_PATH)
